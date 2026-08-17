"""Assemble the final report: insert Chapter 8 and the figures, then renumber.

    python build_report.py [output.docx]

Writes a NEW file; the source document is never modified.

Three things happen, in order:

  1. Figures and Chapter 8 are inserted at NAMED ANCHORS, not appended. A
     results chapter after the conclusion, and figures after the text that
     discusses them, both read wrong.
  2. Every figure and table caption in the document is renumbered in document
     order, chapter by chapter. Inserting anything in the middle otherwise
     leaves the existing captions wrong, and the cross-references with them.
  3. The List of Figures and List of Tables are rebuilt from the renumbered
     captions, with real page numbers read back from a rendered PDF.

Chapter numbers come from Word's own heading numbering, so inserting Chapter 8
before Threats renumbers Threats to 9 and Conclusion to 10 automatically; this
script only has to keep the captions in step.
"""
import os
import re
import shutil
import subprocess
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = r"C:\Users\yc199\Downloads\MIS41170-Capstone.docx"
DST = sys.argv[1] if len(sys.argv) > 1 else \
    r"D:\MSBA\Capstone\TSGP\MIS41170-Capstone-FINAL.docx"
FIG = "figures"

# Front matter and back matter carry no chapter number; the chapters Word
# numbers 1..N run from Introduction to the last numbered chapter.
FIRST_CHAPTER = "Introduction"
AFTER_LAST_CHAPTER = "Appendices"

GREY = RGBColor(0x5F, 0x63, 0x68)

# Page numbers in the two front-matter lists and in the index cannot be known
# until the document is laid out, so every one of them is written first as
# this placeholder -- the width of a real page number, so that filling it in
# afterwards cannot reflow the line it sits on.
PLACEHOLDER_PAGE = "00"
doc = Document(SRC)
STYLES = {s.name for s in doc.styles}
FIGCAP = "Figcaption" if "Figcaption" in STYLES else None
TABCAP = "Tablecaption" if "Tablecaption" in STYLES else None
HEAD_STYLE = "tablehead1" if "tablehead1" in STYLES else None
BODY_STYLE = "tabletext1" if "tabletext1" in STYLES else None
LIST_STYLE = "Paragraph" if "Paragraph" in STYLES else None


# ---------------------------------------------------------------- helpers
def find_heading(text, level=None):
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        if level and not p.style.name.endswith(str(level)):
            continue
        if p.text.strip().lower().startswith(text.lower()):
            return p
    raise LookupError(f"anchor not found: {text!r}")


def after_paras(heading, n):
    """The anchor that places a figure after the first n body paragraphs of a
    section, so the reader meets the diagram once the section has introduced
    it rather than at the tail of the section before.

    A caption never counts as a body paragraph and is never inserted in front
    of: a table and its caption have to stay together, so an anchor landing on
    one advances past it.
    """
    caps = {s for s in (FIGCAP, TABCAP) if s}
    paras = doc.paragraphs
    i = next(k for k, p in enumerate(paras) if p._p is heading._p)
    seen = 0
    for k in range(i + 1, len(paras)):
        p = paras[k]
        if p.style.name.startswith("Heading"):
            return p
        if not p.text.strip() or p.style.name in caps:
            continue
        seen += 1
        if seen >= n:
            for j in range(k + 1, len(paras)):
                q = paras[j]
                if q.text.strip() and q.style.name not in caps:
                    return q
            return paras[-1]
    return paras[-1]


# body text is the document's own "Paragraph" style (Normal plus 9 pt before);
# leaving new text on Normal loses the space between paragraphs
BODY = LIST_STYLE


def para_before(anchor, text="", *, style=None, align=None):
    p = anchor.insert_paragraph_before("", style=style or BODY)
    if align is not None:
        p.alignment = align
    if text:
        p.add_run(text)
    return p


def rich_before(anchor, parts):
    p = anchor.insert_paragraph_before("", style=BODY)
    for t, b in parts:
        r = p.add_run(t)
        r.bold = b
    return p


def unnumber(p):
    """Suppress the automatic heading number on one paragraph.

    Word numbers Heading 1 and 2 throughout, so an appendix heading inserted
    with those styles comes out as a section of the last chapter -- "10.5
    Appendix B". The back-matter headings in the source document suppress it
    the standard way, with numId 0, and this does the same.
    """
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    for tag, val in (("w:ilvl", "0"), ("w:numId", "0")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        numPr.append(el)
    pPr.append(numPr)
    return p


def heading_before(anchor, text, level):
    p = anchor.insert_paragraph_before("", style=f"Heading {level}")
    p.add_run(text)
    return p
    return p


# the text block is narrower than the page suggests (1.575" gutter for
# binding), so anything wider than this runs past the right margin
_s = doc.sections[0]
TEXT_WIDTH = (_s.page_width - _s.left_margin - _s.right_margin) / 914400


def figure_before(anchor, filename, caption_text, width=6.0):
    path = os.path.join(FIG, filename)
    if not os.path.exists(path):
        print(f"  !! missing figure {path}")
        return
    width = min(width, TEXT_WIDTH)
    p = anchor.insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    c = anchor.insert_paragraph_before("", style=FIGCAP)
    c.add_run(caption_text)
    print(f"  + {filename}")


def replace_figure(caption_contains, filename, width=6.0):
    """Swap the image above a caption for a redrawn one, caption untouched."""
    path = os.path.join(FIG, filename)
    if not os.path.exists(path):
        print(f"  !! missing figure {path}")
        return
    paras = doc.paragraphs
    for k, p in enumerate(paras):
        if p.style.name != FIGCAP or caption_contains.lower() not in \
                p.text.lower():
            continue
        for j in range(k - 1, max(k - 6, -1), -1):
            q = paras[j]
            if not q._p.findall(f".//{qn('w:drawing')}"):
                continue
            for r in list(q.runs):
                r._r.getparent().remove(r._r)
            q.alignment = WD_ALIGN_PARAGRAPH.CENTER
            q.add_run().add_picture(path, width=Inches(min(width, TEXT_WIDTH)))
            print(f"  ~ {filename} (replaces the drawn-in version)")
            return
    print(f"  !! no image found above a caption matching {caption_contains!r}")


def shade(cell, fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def set_borders(table, colour="BFBFBF", sz=4):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), colour)
        borders.append(e)
    table._tbl.tblPr.append(borders)


def keep_together(table):
    """Stop a table breaking across a page boundary.

    Without this a header row can be left stranded at the foot of a page with
    its body overleaf. Rows are marked non-splitting, every row but the last
    is kept with the row after it, and the header row is marked as one so it
    repeats if the table ever does have to break.
    """
    hdr = table.rows[0]._tr.get_or_add_trPr()
    hdr.append(OxmlElement("w:tblHeader"))
    for n, row in enumerate(table.rows):
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        if n == len(table.rows) - 1:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True


def table_before(anchor, headers, rows, widths, *, bold_cols=(), left_cols=()):
    def al(i):
        return (WD_ALIGN_PARAGRAPH.LEFT if (i == 0 or i in left_cols)
                else WD_ALIGN_PARAGRAPH.CENTER)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        if HEAD_STYLE:
            p.style = doc.styles[HEAD_STYLE]
        p.alignment = al(i)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        shade(c, "ECEFF1")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            if BODY_STYLE:
                p.style = doc.styles[BODY_STYLE]
            p.alignment = al(i)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(v))
            r.font.size = Pt(9)
            if i in bold_cols:
                r.bold = True
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = Inches(widths[i])
    keep_together(t)
    anchor._p.addprevious(t._tbl)
    return t


def caption_before(anchor, text):
    c = anchor.insert_paragraph_before("", style=TABCAP)
    c.add_run(text)
    return c


# ========================================================= 0. corrections ==
# A fact-check of the written chapters against the archived run files found
# figures carried over from superseded grids, a diagnosis withdrawn in Chapter
# 7 but still asserted in Chapter 1, and cross-references left behind by the
# insertion of the classification chapter. Each is rewritten here rather than
# in the source document, so the source stays the single hand-edited original
# and every correction is visible in one place.
#
# A pattern that fails to match, or matches more often than stated, aborts the
# build: a correction that silently did nothing is worse than none at all.


def _splice(p, old, new):
    """Replace `old` with `new` inside one paragraph, run formatting intact.

    A sentence is normally spread over several runs split at arbitrary points,
    so neither `p.text = ...` (which would flatten the formatting) nor a
    per-run search (which would miss anything spanning a boundary) will do.
    The replacement is written into the run where the match begins, inheriting
    that run's formatting, and the rest of the matched span is cleared; text
    either side of the match is untouched.
    """
    runs = p.runs
    i = "".join(r.text for r in runs).find(old)
    if i < 0:
        return False
    j, pos = i + len(old), 0
    for r in runs:
        s, e = pos, pos + len(r.text)
        pos = e
        if e <= i or s >= j:                      # run lies outside the match
            continue
        head = r.text[:i - s] if s < i else ""
        tail = r.text[j - s:] if e > j else ""
        r.text = head + (new if s <= i else "") + tail
    return True


def correct(old, new, expect=1):
    # A merged cell is handed back once per column it spans, so its paragraphs
    # have to be visited only once. The element is kept as the dict's value,
    # not just its id: lxml builds proxies on demand and frees them when the
    # last reference goes, and a freed id is reused, which had a table cell
    # skipped because a body paragraph happened to have been given its number.
    seen, hits = {}, 0

    def walk(paras):
        nonlocal hits
        for p in paras:
            if id(p._p) in seen:
                continue
            seen[id(p._p)] = p._p
            if _splice(p, old, new):
                hits += 1

    walk(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                walk(cell.paragraphs)
    if hits != expect:
        raise SystemExit(f"CORRECTION matched {hits}x, expected {expect}x:\n"
                         f"  {old[:90]!r}")
    return hits


print("applying corrections to the written chapters ...")

# -- claims the archived results do not support ---------------------------
# RQ1: TSGP has the worse median on 5/5 and is significantly worse on 3/5.
# "four" matched neither count and contradicted Chapters 1 and 6.
correct("TSGP loses to standard genetic programming on four of the five "
        "datasets, and the direction of the comparison agrees with the paper "
        "on one.",
        "TSGP has the worse median test error on all five datasets, "
        "significantly so on three, and the direction of the comparison "
        "agrees with the paper on one.")

# The contributions list still carried the token-accuracy diagnosis that
# Table 7.3 withdraws.
correct("on an intrinsically one-to-many transformation task its residual "
        "token-level error places offspring far outside the semantic distance "
        "it was trained to produce.",
        "its semantic step has a floor well above the distance the method asks "
        "for, and, unlike standard genetic programming, it cannot refine that "
        "step as the population converges.")

# Only Galaxy changed verdict. ESL was already not significant on the old
# grid (p = 0.91), so "two datasets" overstated the effect of the split fix.
correct("Two datasets that were significant losses under the previous "
        "configuration, Galaxy in particular, are no longer distinguishable "
        "once each run draws its own split.",
        "Galaxy, a significant loss under the previous configuration, is no "
        "longer distinguishable once each run draws its own split.")

# The paper's own TSGP-vs-stdGP difference is significant on three of the
# four it leads on; the other two are directions of the medians only.
correct("the reverse of the published finding, in which TSGP wins on four of "
        "the five.",
        "the reverse of the published finding, in which TSGP has the better "
        "median on four of the five, significantly so on three.")
correct("where the paper reports it winning on four.",
        "where the paper reports it ahead on four.")

# Recounted on the final grid: last improving generation 10-29, improving
# generations 3-5. The published pair came from the superseded v1 grid.
correct("stops improving at a median generation between 15 and 32 of the 50, "
        "and only 2.5 to 6 generations in a typical run produce any "
        "improvement at all",
        "stops improving at a median generation between 10 and 29 of the 50, "
        "and only 3 to 5 generations in a typical run produce any improvement "
        "at all")

# diagnostics/baseline_adamw.json, the shipped measurement, gives 0.81/1.67/
# 6.12; the locality ratio in the row below it comes from that same file.
correct("achieved distances of 1.15, 1.55 and 6.44, monotonic across a "
        "5.6-fold span",
        "achieved distances of 0.81, 1.67 and 6.12, monotonic across a "
        "7.6-fold span")

# Division appears in 3-37% of TSGP's best solutions on the final grid (ESL
# 37%); the 0-17% range was measured on the superseded one.
correct("The trained operator almost never emits protected division: it "
        "accounts for 0.7 per cent of tokens in the training targets against "
        "roughly 16 per cent each for addition, subtraction and "
        "multiplication, and it appears in between 0 and 17 per cent of "
        "TSGP's best solutions",
        "The trained operator emits protected division far less readily than "
        "the primitive set allows: it accounts for 0.7 per cent of tokens in "
        "the training targets against roughly 16 per cent each for addition, "
        "subtraction and multiplication, and it appears in 3 to 37 per cent of "
        "TSGP's best solutions")

# Five diagnostic variants were trained: AdamW, batch-32, low-SD, normalised
# SD and binned SD.
correct("the diagnostic investigation reported in Chapter 7 trained four "
        "model variants",
        "the diagnostic investigation reported in Chapter 7 trained five model "
        "variants")

# 103 nodes at 0.443 is a single instrumented run; the medians are 27 and
# 0.453, which makes the same point without reading as a summary.
correct("on ESL it reaches 103 nodes at a test error of 0.443 against a "
        "published 12 nodes at 0.502.",
        "on ESL its median best solution reaches 27 nodes at a test error of "
        "0.453 against a published 12 nodes at 0.502, and individual runs grow "
        "past a hundred nodes.")

# -- what the numbers are, and are not ------------------------------------
# Table 6.3 and the step figure are one instrumented run each; the figures
# either side of them are medians over thirty. Saying so once prevents the
# table being read as a 30-run summary.
correct("Table 6.3 shows all three measured on ESL, where the magnitude of "
        "the target vector is 15.6.",
        "Table 6.3 shows all three measured on a single instrumented run of "
        "each method on ESL, under identical seeds, where the magnitude of the "
        "target vector is 15.6.")
correct("Table 6.3  Per-generation behaviour on ESL (target-vector magnitude "
        "15.6). Step distance is the semantic distance between a parent and "
        "its offspring.",
        "Table 6.3  Per-generation behaviour on ESL (target-vector magnitude "
        "15.6), from one instrumented run of each method under identical "
        "seeds; the figures either side of it are medians over all thirty "
        "runs. Step distance is the semantic distance between a parent and its "
        "offspring.")

# The paper's Galaxy and pollen comparisons are not significant, so the
# paper column there records the direction of the medians.
correct("Table 6.2  Wilcoxon rank-sum comparison of TSGP and standard GP at "
        "the 5% level. The direction agrees with the paper on one of five "
        "datasets, pollen.",
        "Table 6.2  Wilcoxon rank-sum comparison of TSGP and standard GP at "
        "the 5% level. The direction agrees with the paper on one of five "
        "datasets, pollen. The paper's own comparison is significant on ERA, "
        "ESL and LEV only, so its column on Galaxy and pollen records the "
        "direction of the medians rather than a significant result.")

# The printed p-values are rank-sum, not paired; and only Galaxy beats the
# published TSGP figure. The body's claim, against our own k=1 arm, holds.
correct("The modified operator beats the paper's own TSGP on all five "
        "datasets; p is the paired comparison against standard GP.",
        "The modified operator is significantly better than the paper's own "
        "operator, the k=1 column, on all five datasets, and beats the "
        "published TSGP figure on Galaxy; p is the Wilcoxon rank-sum "
        "comparison against standard GP.")

# The target decays and the schedule was tuned, both of which the next
# section's caveat depends on.
correct("keep the one whose semantic distance from the parent is closest to a "
        "target, rather than accepting the first draw.",
        "keep the one whose semantic distance from the parent is closest to a "
        "target, rather than accepting the first draw. The target is expressed "
        "as a fraction of the magnitude of the training target vector and "
        "decays across the run, from 1.0 of it to 0.02; that schedule was "
        "tuned on ESL over ten runs, a caveat the next section returns to.")

# The binned model's floor is at its own level, not at the same value.
correct("and at matched training it exhibits the same floor.",
        "and at matched training it exhibits a floor of the same kind, flat "
        "below a request of about 0.1, though at its own level rather than at "
        "the same value.")

# The floor is the report's central finding, so it is worth saying at the
# point it is established that it does not depend on the training corpus
# either. Chapter 8's operator is the control: different problems, a pool
# built after the Section 5.5 corrections, a scale some forty times larger.
correct("the achieved distance does rise with the request over the range the "
        "model can reach, but it cannot be pushed below its floor.",
        "the achieved distance does rise with the request over the range the "
        "model can reach, but it cannot be pushed below its floor. Nor is it a "
        "defect of the training pool. The operator built for the extension in "
        "Chapter 8 is trained from scratch on pairs drawn from different "
        "problems, generated after both of the data-generation corrections "
        "described in Section 5.5, and at a semantic scale some forty times "
        "larger; it reproduces the same floor. The limit therefore sits in the "
        "operator rather than in any particular corpus it was fitted to.")

# Both references point at the step-size figure, which the size-across-
# generations figure inserted above it pushed from 6.2 to 6.3.
correct("That single contrast, shown in Figure 6.2, explains",
        "That single contrast, shown in Figure 6.3, explains")
correct("over a run on ESL, as Table 6.3 and Figure 6.2 show.",
        "over a run on ESL, as Table 6.3 and Figure 6.3 show.")

# Section 7.8 was quoted from a screening run whose per-run files had not
# been kept. It has been rerun as a proper arm: 30 runs on each dataset,
# drawing the same splits from the same seeds as the baseline so the two pair
# run for run (results_nodiv/, TSGP_NO_DIVISION=1). The direction survives on
# four of five, the magnitudes do not, and nothing is significant.
correct("Testing it first showed the opposite. Removing division from "
        "standard GP's primitive set makes standard GP better on four of five "
        "datasets, by 0.026 on ESL, 0.033 on Galaxy, 0.027 on LEV and 0.039 on "
        "pollen. The operator's reluctance to emit division is thus an "
        "advantage it is already receiving, not a handicap. This carries an "
        "uncomfortable implication worth stating rather than omitting: a "
        "standard GP without division reaches 0.417 on ESL and 0.307 on "
        "Galaxy, beating both the published standard GP figures and our own "
        "best TSGP, so part of TSGP's apparent parity is measured against a "
        "baseline handicapped by a primitive the original paper specifies.",
        "Testing it first showed the opposite. Standard genetic programming "
        "was rerun with division removed from its primitive set, thirty runs "
        "on each dataset drawing their splits from the same seeds as the "
        "baseline, so that the two grids pair run for run. Removing the "
        "primitive leaves standard GP nominally better on four of the five "
        "datasets, by 0.037 on ESL, 0.037 on pollen, 0.014 on Galaxy and 0.001 "
        "on ERA, and worse by 0.003 on LEV, but a paired Wilcoxon test makes "
        "none of the five significant at the 5% level, the closest being ESL "
        "at p = 0.053. The finding is therefore the null one, and it is the "
        "null result that answers the question: if a primitive can be taken "
        "away from standard GP without measurable cost, an operator that "
        "rarely emits it is not thereby handicapped, and regenerating five "
        "million training pairs to encourage it is not a lead worth the ten "
        "hours of regeneration and retraining it would take. One consequence "
        "is worth stating "
        "rather than omitting: standard GP without division reaches 0.416 on "
        "ESL and 0.318 on Galaxy, better on both than the standard GP figures "
        "the paper publishes, so the baseline TSGP is measured against in "
        "Table 6.1 is not one that protected division is holding back.")

# -- corrections that are in the code but not in the results --------------
# DATAGEN_GP_GENERATIONS is 50, and Section 5.5 reports the fix; two passages
# still described the old value as current.
correct("is run with a large population (two thousand individuals) for one "
        "hundred generations.",
        "is run with a large population (two thousand individuals) for fifty "
        "generations, the run length Table 1 of the paper specifies.")
correct(" One outright conflict is also noted: the data-generation code runs "
        "for one hundred generations, whereas Table 1 specifies fifty; the "
        "paper's phrase about generating a sufficient number of functions "
        "gives some licence, but the point is left explicit rather than "
        "hidden.", "")

# The training pairs, and so the operator behind every result here, predate
# both data-generation fixes. Reporting them as closed overstated it.
correct("The pool is now shuffled before pairing.",
        "The pool is now shuffled before pairing. Both defects were found "
        "after the model had been trained, so the correction is in the "
        "generator and not in the pairs the reported operator learned from. "
        "Carrying it through to the results means regenerating five million "
        "pairs, about three hours, retraining for eight epochs at the measured "
        "forty-nine minutes an epoch, and then repeating every measurement "
        "downstream of the model: the thirty-run grid, the step-control grid "
        "at ten and a half hours, and the whole of the Chapter 7 diagnostic "
        "sequence. That is upwards of twenty-two hours of continuous work on "
        "the single graphics card available, which on this machine had also to "
        "be scheduled around a Windows update service that terminated two "
        "earlier long runs. The budget was spent instead on the measurement "
        "that decides whether it would have mattered, reported in Section 8.6: "
        "a second operator, trained from scratch on a pool built after both "
        "corrections, reproduces the same step-size floor.")
correct("noted as an open conflict in the previous draft, this is now "
        "corrected.",
        "noted as an open conflict in the previous draft, this is now "
        "corrected in the generator, and carries the same caveat and the same "
        "control as the defect above.")
correct("Two deviations recorded in earlier drafts are now closed. Data "
        "generation, which had run for one hundred generations against the "
        "fifty specified in Table 1, has been corrected to fifty; and the "
        "train and test partition, which had been shared across all thirty "
        "runs of a dataset, is now drawn independently for each run. Two "
        "departures remain in the trained model,",
        "One deviation recorded in earlier drafts is fully closed: the train "
        "and test partition, which had been shared across all thirty runs of a "
        "dataset, is now drawn independently for each run, and the whole grid "
        "was rerun on that basis. A second is closed in the generator but not "
        "in the trained model. Data generation, which had run for one hundred "
        "generations against the fifty specified in Table 1, has been "
        "corrected to fifty, and the pair-selection defect described in "
        "Section 5.5 has been fixed, but both were found after training, and "
        "carrying them through would have cost the twenty-two hours of "
        "regeneration, retraining and re-measurement itemised in that section. "
        "The alternative was to measure whether it would change the finding, "
        "which is the cheaper and the more informative of the two: the "
        "classification operator of Chapter 8 is trained from scratch on a "
        "pool built after both corrections, on a different task and at a "
        "semantic scale some forty times larger, and it reproduces the "
        "step-size floor unchanged. The diagnosis this report rests on is "
        "therefore established on a corrected pool as well as on the original "
        "one. Two further departures remain in the trained model,")

# -- the datasets ---------------------------------------------------------
# The paper describes four real-world problems and one synthetic (Pollen).
correct("on five real-world benchmark datasets.", "on five benchmark datasets.")
correct("across five real-world symbolic regression datasets,",
        "across five symbolic regression benchmark datasets,")
correct("across five real-world benchmark datasets, while keeping solution "
        "sizes small.",
        "across five benchmark datasets, while keeping solution sizes small.")
correct("Evaluation uses five real-world regression datasets drawn from",
        "Evaluation uses five regression datasets drawn from")
correct("so that the comparison is made on the same ground. They are "
        "black-box problems in the sense that no closed-form generating "
        "function is known,",
        "so that the comparison is made on the same ground; four are "
        "real-world problems and pollen is synthetic, following the original's "
        "own description of them. All five are treated as black-box problems, "
        "in that no generating function is assumed or used,")

# Every one of the five is natively four-dimensional, so nothing is
# discarded: the constraint bears on which problems are usable at all.
correct("A deliberate restriction applies to all five: only the first four "
        "features of each dataset are used. This matches the four-dimensional "
        "setting of the original study and is required by the method itself, "
        "because the primitive terminal set is fixed at four input variables "
        "(x0 to x3) and the transformer is trained only on expressions over "
        "those variables. Using the first four features is therefore not an "
        "incidental preprocessing choice but a structural constraint of the "
        "trained operator, and its implications for what the results can claim "
        "are noted in Chapter 8.",
        "A structural constraint governs which datasets can be used at all: "
        "the primitive terminal set is fixed at four input variables (x0 to "
        "x3) and the transformer is trained only on expressions over those "
        "variables, so a benchmark must be expressible in four features. All "
        "five are natively four-dimensional, exactly as in the original study, "
        "so no feature is discarded here; the loader takes the first four as a "
        "safeguard rather than as a reduction. The constraint therefore bears "
        "on the range of problems the trained operator can address at all, not "
        "on the preprocessing of these five, and its implications for what the "
        "results can claim are noted in Chapter 9.")
correct("in every case the first four features are used, matching the "
        "four-dimensional problem setting adopted throughout.",
        "in every case the dataset is natively four-dimensional, matching the "
        "problem setting adopted throughout.")

# -- cross-references the new chapter renumbers ---------------------------
# Word numbers the headings, so inserting Chapter 8 pushes Threats to 9 and
# Conclusion to 10; the references written by hand do not follow.
correct("and Chapters 6 to 8 present the evidence and its interpretation in "
        "full.",
        "and Chapters 6, 7 and 9 present the evidence and its interpretation "
        "in full.")
correct("Chapter 8 sets out the threats to validity and the known deviations "
        "from the original method, and Chapter 9 concludes and outlines future "
        "work.",
        "Chapter 8 extends the method to binary classification, a task the "
        "original paper does not address. Chapter 9 sets out the threats to "
        "validity and the known deviations from the original method, and "
        "Chapter 10 concludes and outlines future work.")
correct("and its consequences for what can and cannot be concluded, is "
        "discussed in Chapter 8.",
        "and its consequences for what can and cannot be concluded, is "
        "discussed in Chapter 9.")

# The third future-work item was written before Chapter 8 measured the
# semantic regime and retrained on a corrected pool, which is most of what it
# asks for; what is left is the narrower question.
correct("The third is to validate the training pool, on which every hypothesis "
        "tested here is downstream and whose construction has not been checked "
        "against the semantic regime the search actually occupies.",
        "The third concerns the training pool, on which every hypothesis tested "
        "here is downstream. Chapter 8 measures the semantic regime a search "
        "actually occupies against the regime the pool covers, and retrains the "
        "operator on a pool built to match it, which settles the question for "
        "the step-size floor; what remains is to build a regression pool "
        "matched in the same way and ask whether anything other than the floor "
        "responds to it.")

# -- proofreading ---------------------------------------------------------
# The report has sections 3.2, 4.1 and 4.2 of its own, so a bare reference to
# the paper's sections of those numbers sends a reader to the wrong document.
correct("Section 4.1 specifies a 10% terminal bias in subtree crossover.",
        "Section 4.1 of the paper specifies a 10% terminal bias in subtree "
        "crossover.")
correct("Section 4.1 specifies AdamW; the original run used Adam.",
        "Section 4.1 of the paper specifies AdamW; the original run used "
        "Adam.")
correct("Section 3.2 describes the desired distance as the means of "
        "controlling step size",
        "Section 3.2 of the paper describes the desired distance as the means "
        "of controlling step size")

# What is returned is the best on the training data; the test set is what it
# is then measured on, and the two must not be run together in one clause.
correct("The process repeats for a fixed number of generations, and the best "
        "individual on unseen data is returned.",
        "The process repeats for a fixed number of generations, and the best "
        "individual found is returned and measured on unseen data.")

# "the authors" a sentence after the paper's authors reads as theirs
correct("It is very recent and, to the best of the authors' knowledge, not "
        "yet independently replicated;",
        "It is very recent and, to the best of our knowledge, not yet "
        "independently replicated;")

# pollen is synthetic, as Section 4.1 now says, so the benchmarks cannot be
# described as real-world wholesale
correct("both the synthetic problems used to build the model and the "
        "real-world benchmarks used to evaluate it.",
        "both the synthetic problems used to build the model and the "
        "benchmarks used to evaluate it.")
correct("a separate body of data is used to evaluate it, namely the "
        "real-world benchmarks on which the evolutionary search is run.",
        "a separate body of data is used to evaluate it, namely the benchmark "
        "datasets on which the evolutionary search is run.")

print("corrections applied")


# ============================================================ 1. insert ===
print("inserting figures into existing chapters ...")

# two figures carried a drawn-in title (now the caption's job) and a printed
# number on every mark, which collided with the tick labels. Both are redrawn
# from the same runs, so the captions and Table 6.1 still hold.
replace_figure("Median test RMSE over 30 runs, ours", "fig_rmse_vs_paper.png")
replace_figure("Parent-to-offspring semantic distance", "fig_semantic_step.png")

# each diagram sits inside the section that introduces it, one paragraph in,
# rather than at the tail of the section before it

# Chapter 2 described the loop and its operators in prose alone. The cycle
# goes after the paragraph that walks through it, and the operators go at the
# head of the semantic section, where their blindness to behaviour is the
# argument being made.
a = after_paras(find_heading("Genetic programming for symbolic regression",
                             level=2), 2)
figure_before(a, "fig_gp_cycle.png",
              "PLACEHOLDER  The genetic programming loop. Each generation "
              "evaluates the population, selects parents and produces the next "
              "population by crossover and mutation, repeating until the "
              "generation limit is reached. TSGP replaces one box of this "
              "diagram, the variation step, and leaves the rest of the loop as "
              "it stands.", width=6.0)

a = after_paras(find_heading("Semantic and geometric semantic genetic "
                             "programming", level=2), 1)
figure_before(a, "fig_gp_operators.png",
              "PLACEHOLDER  Subtree crossover exchanges a branch between two "
              "parents; subtree mutation replaces a branch with a randomly "
              "generated one. Both act on structure alone, with no regard for "
              "what the expression computes, which is the limitation the "
              "semantic methods address.", width=5.8)

a = after_paras(find_heading("The transformer architecture"), 1)
figure_before(a, "fig_transformer_architecture.png",
              "PLACEHOLDER  The encoder-decoder transformer of Vaswani et al. "
              "(2017). The encoder maps the input sequence to a representation; "
              "the decoder generates the output one token at a time, attending "
              "both to what it has produced so far and, through cross-attention, "
              "to the encoder's representation. Both stacks repeat the same "
              "block N times.", width=5.3)

a = after_paras(find_heading("Semantics and semantic distance", level=2), 1)
figure_before(a, "fig_semantic_distance.png",
              "PLACEHOLDER  Semantics and semantic distance. The semantics of "
              "an expression is the vector of values it produces on a fixed set "
              "of probe inputs; the semantic distance between two expressions "
              "is the Euclidean distance between those vectors. Two expressions "
              "are semantically similar when the vectors are close, however "
              "differently the expressions are built.", width=6.2)

# the model diagram belongs with our implementation (Methodology, Phase 2),
# not with the literature description of the method
a = after_paras(find_heading("Phase 2"), 1)
figure_before(a, "fig_tsgp_model.png",
              "PLACEHOLDER  The model used in this study. A parent expression "
              "is tokenised in prefix order and encoded together with the "
              "desired semantic distance; the decoder then emits an offspring "
              "one token at a time, with syntax control masking any token that "
              "would make a valid tree impossible.", width=6.2)

a = after_paras(find_heading("Solution size", level=2), 1)
figure_before(a, "fig_size_generations.png",
              "PLACEHOLDER  Size of the best solution across a run on ESL, "
              "median and interquartile range over 30 runs. Standard GP "
              "continues to build structure for as long as it keeps finding "
              "improvements; TSGP stops early and stays small.", width=5.2)

a = after_paras(find_heading("The operator cannot take a step", level=2), 1)
figure_before(a, "fig_step_floor.png",
              "PLACEHOLDER  Achieved parent-to-offspring distance against the "
              "distance requested, both on logarithmic scales. Below a request "
              "of roughly 0.1 the achieved distance does not fall further. The "
              "dashed line shows where the points would lie if the request were "
              "honoured; the method's operating point lies beneath the floor.",
              width=5.2)

a = after_paras(find_heading("What controlled measurement overturned", level=2), 1)
figure_before(a, "fig_locality_control.png",
              "PLACEHOLDER  The control condition. Distance from an offspring "
              "to its own parent, against distance from the same offspring to "
              "an unrelated parent, at each requested distance. The ratio above "
              "each pair is the first divided by the second; a value near 1 "
              "would mean the parent had been ignored.", width=5.2)

print("inserting Chapter 8 ...")
A = find_heading("Threats to Validity", level=1)

heading_before(A, "Extending TSGP to Binary Classification", 1)
para_before(A, "The preceding chapters test a published claim and find that it "
               "does not reproduce as specified. This chapter asks a different "
               "question. The mechanism at the centre of TSGP, a transformer "
               "that proposes structurally free but semantically close "
               "variations, is not specific to regression. If it has value, "
               "that value should show somewhere. This chapter extends the "
               "method to binary classification, a task the original paper does "
               "not address, and reports what it finds.")
rich_before(A, [
    ("The extension is not a second replication. ", True),
    ("It is an original application of the method, and its results should be "
     "read as such: they say nothing about whether the published regression "
     "result is correct, and they are not evidence for or against the diagnosis "
     "of the preceding chapter.", False)])

heading_before(A, "Keeping semantics continuous: decision values instead of "
                  "labels", 2)
para_before(A, "The obvious way to apply a symbolic method to classification is "
               "to have the evolved expression emit a class label. That would "
               "break TSGP specifically, and it is worth being precise about "
               "why.")
para_before(A, "The method's two load-bearing definitions are that the "
               "semantics of an expression is the vector of values it produces "
               "on a set of probe inputs, and that the semantic distance "
               "between two expressions is the Euclidean distance between those "
               "vectors. An expression that emits only zeros and ones has a "
               "semantics that is a binary vector, and the distance between two "
               "such vectors is simply a count of the samples they disagree on. "
               "That measure is coarse and largely flat: very many structurally "
               "unrelated expressions sit at exactly the same distance from one "
               "another. The k-nearest-neighbour search that manufactures the "
               "transformer's training pairs would then be matching "
               "near-arbitrary expressions, and the operator would have nothing "
               "to learn.")
para_before(A, "The expression is therefore kept real-valued and its output is "
               "reinterpreted. The predicted class is the sign of the output, "
               "with labels encoded as -1 and +1; the surface on which the "
               "expression evaluates to zero is the decision boundary, and the "
               "magnitude of the output is a measure of confidence. This is the "
               "same device that underlies logistic regression and support "
               "vector machines. Its value here is specific: semantics remain "
               "continuous vectors, Euclidean distance retains its meaning, and "
               "the entire transformer apparatus transfers without modification.")
figure_before(A, "fig_decision_value.png",
              "PLACEHOLDER  A real-valued expression read as a classifier. The "
              "sign of the output gives the class and its magnitude gives "
              "confidence, so the semantics of the expression remain a "
              "continuous vector and semantic distance keeps its meaning.",
              width=6.0)

heading_before(A, "Fitness: a smooth surrogate rather than accuracy", 2)
para_before(A, "Selection cannot act on accuracy directly. Accuracy is "
               "piecewise constant in the parameters of an expression: a small "
               "change to a coefficient leaves it completely unchanged until "
               "some sample crosses the boundary, at which point it jumps. "
               "Selection would be climbing a staircase, with no gradient "
               "between steps.")
para_before(A, "Fitness is therefore the mean logistic loss of the decision "
               "value, log(1 + exp(-y f(x))), where the quantity y f(x) is the "
               "classification margin: positive when the prediction is correct "
               "and larger when it is more confidently correct. This is smooth, "
               "so every small improvement registers. Accuracy and the area "
               "under the ROC curve are computed for reporting only and never "
               "used for selection.")

heading_before(A, "Constructing the benchmarks", 2)
para_before(A, "The benchmark set had to be constructed, and the reason should "
               "be stated plainly. The trained operator's terminal set is fixed "
               "at four variables, so a problem must be expressible in four "
               "features. The Penn Machine Learning Benchmarks collection "
               "contains exactly two four-feature binary classification "
               "datasets, and neither is usable for a benchmark: one has 264 "
               "samples at 73% class imbalance, the other 50 samples.")
para_before(A, "The five regression benchmarks are therefore converted into "
               "binary problems by thresholding the target. The features are "
               "the benchmarks' own, unchanged, every problem is natively "
               "four-dimensional so nothing is discarded, and the "
               "classification and regression results are measured on identical "
               "data, which allows the operator's behaviour to be compared "
               "across the two tasks rather than across different benchmarks. "
               "Two constructions are used, and the second exists for a reason "
               "established in the next section.")
table_before(A, ["Construction", "Definition", "Rationale"],
             [["Median split",
               "y above a cut chosen from the observed values to divide the "
               "training half as close to evenly as possible",
               "The natural construction. A plain median is insufficient "
               "because these targets are discrete ordinal ratings with heavy "
               "ties: on LEV it produces a 21/79 split whose majority baseline "
               "would swamp any signal."],
              ["Middle band",
               "y inside the central third of the training distribution",
               "Not a linear function of the features even when y is, so a "
               "linear model cannot represent it. Introduced after the analysis "
               "in the next section."]],
             [1.05, 2.15, 3.0], left_cols=(1, 2))
caption_before(A, "PLACEHOLDER  The two label constructions. Every result file "
                  "records which construction produced it.")

heading_before(A, "Establishing that the benchmarks discriminate", 2)
para_before(A, "Comparing a new method only against the baseline it is designed "
               "to improve on is a common way to obtain a flattering result. "
               "Standard classifiers were therefore run on the same splits, and "
               "the outcome changed the design of the study.")
rich_before(A, [
    ("On the median-split task, logistic regression with five coefficients "
     "matches or beats TSGP on all five datasets", True),
    (", significantly on ERA and LEV and statistically indistinguishably on the "
     "other three, and a random forest performs worse than logistic regression "
     "on all five of "
     "them. That pattern is the signature of a decision boundary that is close "
     "to linear, and it follows directly from the construction: thresholding a "
     "monotonic target produces a boundary a linear model captures almost "
     "exactly. The median-split task therefore cannot test whether structural "
     "flexibility is worth anything, because there is no non-linearity for "
     "flexibility to exploit.", False)])
para_before(A, "The middle-band construction was introduced to address this, "
               "and was selected on the evidence of the standard classifiers "
               "alone, before any TSGP run was performed on it. On that task "
               "logistic regression falls to the majority baseline on ERA and "
               "on pollen, while a random forest reaches 0.8525 on ESL against "
               "logistic regression's 0.6434. There is genuine non-linear "
               "structure present that a linear model cannot express.")
figure_before(A, "fig_task_difficulty.png",
              "PLACEHOLDER  Accuracy above the majority-class baseline for "
              "three standard classifiers, on each construction. On the median "
              "split a linear model is as good as a forest; on the middle band "
              "it collapses while the tree-based models hold up. This is why "
              "the second construction exists.", width=6.3)

heading_before(A, "Results at equal budget", 2)
para_before(A, "All results below draw one offspring per parent, as the paper "
               "specifies, so that TSGP and standard GP consume the same number "
               "of model evaluations. Thirty independent runs per method per "
               "dataset, as elsewhere in this report. The step-control variant "
               "is discussed at the end of the section.")
para_before(A, "One parameter has to be rescaled rather than carried over. "
               "Semantic distance is an absolute Euclidean norm, so its scale "
               "follows the magnitude of the functions being compared, and "
               "classification functions are far larger than regression ones "
               "because logistic loss never pins their magnitude down. Measured "
               "on the two training pools, the median distance between paired "
               "expressions is 0.164 for regression and 6.637 for "
               "classification. The paper's requested distance of 0.1 therefore "
               "sits near the middle of the regression pairs but below the "
               "first percentile of the classification ones, which would query "
               "the operator at a distance it effectively never saw in "
               "training. It is replaced throughout this chapter by 2.0, the "
               "lower quartile of the classification pool and so the same "
               "position within it that 0.1 occupies for regression. Both TSGP "
               "arms, including the transfer control of the final section, use "
               "that value, so the comparisons below are internally consistent; "
               "it is a deviation from the paper's stated value and is recorded "
               "as one.")
table_before(A, ["Dataset", "TSGP acc", "stdGP acc", "p", "TSGP size",
                 "stdGP size", "ratio", "p (size)"],
             [["ERA", "0.6870", "0.7020", "0.0015", "9", "31", "3.4x", "7.4e-05"],
              ["ESL", "0.9324", "0.9344", "0.918", "29", "78", "2.7x", "6.7e-10"],
              ["Galaxy", "0.9444", "0.9444", "0.842", "27", "137", "5.1x", "7.4e-11"],
              ["LEV", "0.8060", "0.8330", "<1e-4", "15", "27", "1.8x", "1.2e-04"],
              ["pollen", "0.8381", "0.8363", "0.647", "23", "65", "2.8x", "4.7e-08"]],
             [0.72, 0.82, 0.82, 0.66, 0.82, 0.82, 0.6, 0.84], bold_cols=(4, 6))
caption_before(A, "PLACEHOLDER  Median-split task, equal budget, 30 runs. "
                  "Accuracy compared by Wilcoxon rank-sum.")
table_before(A, ["Dataset", "TSGP acc", "stdGP acc", "p", "TSGP size",
                 "stdGP size", "ratio", "p (size)"],
             [["ERA", "0.5850", "0.5890", "0.762", "9", "79", "8.8x", "3.6e-10"],
              ["ESL", "0.7705", "0.8299", "0.0002", "12", "70", "5.8x", "6.3e-08"],
              ["Galaxy", "0.8704", "0.8735", "0.156", "28", "98", "3.5x", "1.7e-07"],
              ["LEV", "0.6880", "0.6930", "0.264", "10", "21", "2.1x", "1.8e-02"],
              ["pollen", "0.6783", "0.6596", "0.0095", "11", "42", "3.8x", "9.1e-03"]],
             [0.72, 0.82, 0.82, 0.66, 0.82, 0.82, 0.6, 0.84], bold_cols=(4, 6))
caption_before(A, "PLACEHOLDER  Middle-band task, equal budget, 30 runs.")

para_before(A, "Two things stand out, and they should be reported together "
               "because the result is a trade rather than a dominance.")
rich_before(A, [
    ("Solution size. ", True),
    ("TSGP produces significantly smaller solutions on all five datasets in "
     "both tasks, by factors of 1.8 to 8.8, with p-values as low as 7.4e-11. On "
     "Galaxy under the median split it matches standard GP's accuracy exactly "
     "using 27 nodes against 137. This is the paper's central claim about "
     "compactness, holding on a task the paper does not address, and at equal "
     "computational budget.", False)])
rich_before(A, [
    ("Accuracy. ", True),
    ("TSGP is statistically indistinguishable from standard GP on three of five "
     "datasets in each task, and significantly worse on two under the median "
     "split and one under the middle band. It does not match standard GP's "
     "accuracy across the board, and reporting it as though it did would "
     "overstate the result.", False)])
figure_before(A, "fig_size_vs_accuracy.png",
              "PLACEHOLDER  Test accuracy against solution size on the "
              "median-split task, one point per method. TSGP sits far to the "
              "left of standard GP at comparable height: the same accuracy for "
              "a fraction of the structure. Logistic regression, at five "
              "coefficients, is the strongest or joint-strongest model on this "
              "task.", width=6.4)
rich_before(A, [
    ("The exception is instructive. ", True),
    ("On pollen under the middle-band construction TSGP is significantly better "
     "than standard GP (0.6783 against 0.6596, p = 0.0095), and the detail that "
     "gives this meaning is that standard GP there does not beat the "
     "majority-class baseline at all while TSGP does: against the shared "
     "baseline of 0.671 used for the dataset the two sit at -0.011 and +0.007, "
     "and pairing each run with the baseline of its own split leaves standard "
     "GP exactly level with it while TSGP is 0.019 above. The "
     "area under the ROC curve tells the same story more starkly, 0.7421 "
     "against 0.5946 (p < 1e-4). On the hardest non-linear problem in the set, "
     "standard GP finds essentially nothing and the learned operator finds real "
     "structure.", False)])
para_before(A, "The step-control variant was also run on both tasks. It behaves "
               "comparably and wins on pollen under the middle band by a "
               "smaller margin (p = 0.042). Those runs spend eight model "
               "evaluations per offspring and were performed at ten runs rather "
               "than thirty, which is a deliberate allocation rather than an "
               "omission: a step-control unit takes about two minutes against "
               "twenty seconds for an equal-budget one, so raising both arms to "
               "thirty runs would have cost a further seven hours, and it is "
               "the equal-budget arms that every claim in this chapter rests "
               "on. Since ten runs are demonstrably capable of concealing a "
               "genuine difference, as taking the equal-budget arms from ten to "
               "thirty showed twice over, the step-control figures are reported "
               "here as preliminary and nothing is claimed from them.")

heading_before(A, "Whether retraining on classification data was necessary", 2)
para_before(A, "The operator used above was trained on a pool generated "
               "specifically for this task. The reason was measured rather than "
               "assumed: under logistic loss nothing constrains the magnitude "
               "of the expression's output, because a more confident prediction "
               "always scores better, so a classification search drifts towards "
               "ever larger semantic norms. The population's median norm was "
               "observed to climb from 19.5 to 27.3 to 57.7 across generations, "
               "against a regression control that stays near 10, which is where "
               "the regression training pool sits. A pool built from regression "
               "problems therefore leaves a classification search outside the "
               "region the operator was trained on for most of its run.")
para_before(A, "Whether that mattered is a separate question, and it is "
               "answered by running the regression-trained operator on the same "
               "benchmarks, at the same budget and the same number of runs.")
table_before(A, ["Task", "Classification-trained better on", "Significant"],
             [["Median split", "4 of 5",
               "2 of 5: pollen +0.054 (p < 1e-4), ESL +0.010 (p = 0.0065). "
               "Worse on ERA (p = 0.030)."],
              ["Middle band", "3 of 5",
               "1 of 5: pollen +0.026 (p = 0.0004)."]],
             [1.25, 1.85, 3.1], left_cols=(1, 2))
caption_before(A, "PLACEHOLDER  Effect of retraining the operator on "
                  "classification-regime data, both arms at equal budget and "
                  "30 runs.")
rich_before(A, [
    ("The effect is real but modest, and concentrated on a single dataset. ",
     True),
    ("Set against its cost, a complete regeneration of the training pool and "
     "several hours of retraining, this is a negative result worth recording: "
     "an operator trained on regression semantics transfers to classification "
     "better than the semantic-regime measurement predicted it would.", False)])
para_before(A, "One further observation carries across from the preceding "
               "chapter. The step-size floor described there is present here "
               "unchanged: sweeping the requested semantic distance moves the "
               "achieved distance by less than a factor of two. Retraining the "
               "operator on a different task, with a different pool and a "
               "different semantic scale, did not affect it. That strengthens "
               "the reading of the floor as a structural property of the "
               "operator rather than an artefact of any particular training set.")

heading_before(A, "What the extension establishes, and what it does not", 2)
para_before(A, "It establishes that the method transfers. A learned semantic "
               "operator, applied to a task its training never anticipated, "
               "produces solutions 1.8 to 8.8 times more compact than standard "
               "genetic programming at equal budget, with comparable accuracy "
               "on the majority of datasets, and on one non-linear problem "
               "finds structure that standard genetic programming does not find "
               "at all. For an application where a model must be read and "
               "defended as well as scored, that trade is the one worth having.")
para_before(A, "It does not establish that TSGP is a competitive classifier in "
               "general. Logistic regression, at five coefficients, matches or "
               "beats TSGP on every dataset of the median-split task. The "
               "benchmarks are "
               "constructed rather than native, and inherit the four-feature "
               "restriction of the trained operator. The comparison is against "
               "standard genetic programming alone, not against the wider "
               "field. And the accuracy losses on two datasets are real and are "
               "not explained here.")

# ---------------------------------------------- extra threats + conclusion
T = find_heading("Generalisation of the findings", level=2)
heading_before(T, "Limitations of the classification extension", 2)
para_before(T, "The classification benchmarks are constructed by thresholding "
               "the regression targets rather than taken from a native "
               "classification collection, and they inherit the four-feature "
               "restriction of the trained operator. On the median-split "
               "construction logistic regression matches or beats TSGP on every "
               "dataset, which bounds what the extension can claim. The "
               "requested semantic distance is rescaled from the paper's 0.1 to "
               "2.0 for both arms, for the reason given in Section 8.5, so the "
               "chapter reproduces the paper's operator in every respect but "
               "that one. The step-control arms of that chapter were run at ten "
               "repetitions rather than thirty, the seven hours needed to raise "
               "them having been spent instead on bringing every equal-budget "
               "arm to thirty, and they are reported as preliminary with no "
               "claim resting on them. And the comparison is against standard "
               "genetic programming alone.")

print("inserting the end matter ...")


def clear_section(heading_text):
    """Empty a section's body, leaving its heading, and return the heading.

    Only paragraphs and tables are removed, and the walk stops at anything
    else. The body's own sectPr is a sibling of them and is the last element
    of the document, so a loop that deletes whatever it finds takes the
    section properties with it when it clears the final section -- which cost
    this document its binding gutter, its margins and the restart of page
    numbering at Chapter 1, and did so silently, since python-docx writes a
    default sectPr back on save.
    """
    from docx.text.paragraph import Paragraph
    head = find_heading(heading_text, level=1)
    node = head._p.getnext()
    while node is not None and node.tag in (qn("w:p"), qn("w:tbl")):
        if node.tag == qn("w:p") and \
                Paragraph(node, doc).style.name.startswith("Heading"):
            break
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt
    return head


# ---------------------------------------------------------------- appendices
# The stub said the parameter inventory and the commands were in Chapter 5.
# Chapter 5 has the parameters the paper leaves open, not all of them, so the
# inventory is assembled here from config.py instead: one place a reader can
# check every value the study ran with, against what the paper specifies.
A = find_heading("Glossary of terms", level=1)
clear_section("Appendices")
para_before(A, "Three appendices follow. Appendix A is the complete parameter "
               "inventory, marking for each value whether the original paper "
               "specifies it or it was chosen here. Appendix B lists the "
               "commands that reproduce the study end to end. Appendix C gives "
               "the full distribution behind every median quoted in Chapters 6 "
               "to 8, since a median alone says nothing about spread.")

unnumber(heading_before(A, "Appendix A  Parameter inventory", 2))
para_before(A, "Values are as run. The source column distinguishes what the "
               "paper states from what it leaves open, which is the "
               "distinction Chapter 7 turns on: a discrepancy traceable to a "
               "specified value would be an implementation fault, and one "
               "traceable to an unspecified value would not.")
PAPER, OURS = "Paper", "Ours"
table_before(A, ["Parameter", "Value", "Source"],
             [["Population size", "100", PAPER],
              ["Generations", "50", PAPER],
              ["Initialisation", "Ramped half-and-half, depth 2 to 5", PAPER],
              ["Maximum tree depth", "17", PAPER],
              ["Selection", "Tournament, size 5", PAPER],
              ["Function set", "add, sub, mul, protected div", PAPER],
              ["Terminal set", "x0 to x3, ephemeral random constants", PAPER],
              ["ERC grid", "-0.5 to 0.5 in steps of 0.1", PAPER],
              ["Fitness", "Root mean squared error", PAPER],
              ["Crossover", "Subtree, p = 0.9, 10% terminal bias", PAPER],
              ["Mutation", "Subtree, p = 0.1, new subtree depth 0 to 2",
               PAPER],
              ["Runs per dataset", "30", PAPER],
              ["Train and test split", "50/50, both standardised", PAPER],
              ["Replacement", "Generational, no elitism", OURS]],
             [2.0, 2.75, 0.85], left_cols=(1,))
caption_before(A, "Table A.1  Search and evaluation, as used for both TSGP "
                  "and the standard GP baseline.")

table_before(A, ["Parameter", "Value", "Source"],
             [["Synthetic problems", "50", PAPER],
              ["Population, data generation", "2,000", PAPER],
              ["Generations, data generation", "50", PAPER],
              ["Selection, data generation", "Double tournament", PAPER],
              ["Samples per synthetic problem", "200", OURS],
              ["Probe inputs for semantics", "100", OURS],
              ["Neighbours per function, k", "3", PAPER],
              ["Pair filter", "0 < SD < 100", PAPER],
              ["Training pairs", "5,000,000", PAPER],
              ["Encoder and decoder layers", "2 and 2", PAPER],
              ["Attention heads", "8", PAPER],
              ["Hidden dimension", "128", PAPER],
              ["Feed-forward dimension", "512, four times hidden", OURS],
              ["Maximum sequence length", "100 tokens", PAPER],
              ["Vocabulary", "22 tokens", "Follows"],
              ["Parameters", "About 934,000", "Follows"],
              ["Optimiser", "AdamW, learning rate 1e-3", PAPER],
              ["Weight decay", "0.004", OURS],
              ["Epochs", "8", PAPER],
              ["Batch size", "256", OURS],
              ["Loss", "Masked cross-entropy", OURS]],
             [2.0, 2.75, 0.85], left_cols=(1,))
caption_before(A, "Table A.2  Model building: the synthetic corpus in the "
                  "upper block, the transformer in the lower. Follows means "
                  "the value is fixed by others rather than chosen.")

table_before(A, ["Parameter", "Value", "Source"],
             [["Desired distance, regression", "0.1", PAPER],
              ["Desired distance, classification", "2.0, the pool's lower "
               "quartile", OURS],
              ["Sampling temperature", "1.0, no scaling", OURS],
              ["Candidates per parent, k", "1 as the paper specifies; 8 for "
               "the step-control extension", PAPER]],
             [2.0, 2.75, 0.85], left_cols=(1,))
caption_before(A, "Table A.3  The operator as queried during search. Only the "
                  "first row applies to the replication proper; the rest are "
                  "recorded in every result file.")

unnumber(heading_before(A, "Appendix B  Reproducing the study", 2))
para_before(A, "Each stage writes files the next one reads, and every stage "
               "skips work already on disk, so an interrupted run is resumed "
               "by repeating its command. The first four stages build the "
               "operator and take roughly a working day in total; everything "
               "after them reuses it.")
for line, note in (
        ("python -m tsgp.fetch_datasets",
         "caches the five PMLB benchmarks locally, after which no network "
         "access is needed"),
        ("python -m tsgp.data_generator --output data/training",
         "runs standard GP on the 50 synthetic problems and pairs the "
         "expressions it visits by semantic similarity, about three hours"),
        ("python -m tsgp.train_transformer --data data/training "
         "--checkpoints checkpoints_adamw",
         "trains the operator, eight epochs at about 49 minutes each on the "
         "GPU, with a checkpoint after every epoch"),
        ("python -m tsgp.run_experiments --weights "
         "checkpoints_adamw/tsgp_final.npy --output results_v7",
         "the headline grid: both methods, five datasets, 30 runs each"),
        ("python -m tsgp.run_experiments ... --step-k 8 --step-anneal "
         "--step-frac-start 1.0 --step-frac-end 0.02 --output results_anneal_b",
         "the step-control extension of Section 7.7"),
        ("set TSGP_NO_DIVISION=1  &  python -m tsgp.run_experiments "
         "--methods stdgp --output results_nodiv",
         "the primitive-set comparison of Section 7.8"),
        ("run_classification_pipeline.cmd, then run_remaining.cmd, "
         "run_strengthen.cmd and run_final.cmd",
         "the classification chapter in order: pool, training, then each grid"),
        ("python summarise_classification.py  and  python baselines_ml.py",
         "the classification tables and the standard-classifier baselines"),
        ("python make_figures.py  then  python build_report.py",
         "redraws every figure from the result files and rebuilds this "
         "document")):
    # body text is justified, which stretches a monospace command into
    # widely spaced columns; commands are set flush left instead
    p = para_before(A, "", align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(line)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.add_run(f"\n{note}").italic = True

unnumber(heading_before(A, "Appendix C  Full distributions", 2))
para_before(A, "Chapters 6 to 8 quote medians, which is what the original "
               "paper reports and what the rank-sum test compares. The spread "
               "matters too, and in one respect it matters more: the standard "
               "GP baseline is far more variable from run to run than TSGP, so "
               "the two methods differ in consistency as well as in level.")


# The two tables below are read off the result files at build time rather than
# typed in, for the same reason the figures are: a number that has to be
# copied by hand is a number that can fall out of step with the data.
def _runs(directory, method):
    import glob
    import json
    out = {}
    for f in glob.glob(os.path.join(directory, "runs", "*.json")):
        r = json.load(open(f))
        if r["method"] == method:
            out.setdefault(r["dataset"], []).append(r)
    return out


def _stat(rows, key):
    import numpy as np
    v = np.array([r[key] for r in rows], dtype=float)
    return (np.median(v), v.mean(),
            np.percentile(v, 25), np.percentile(v, 75), len(v))


DS = [("1030_ERA", "ERA"), ("1027_ESL", "ESL"),
      ("690_visualizing_galaxy", "Galaxy"), ("1029_LEV", "LEV"),
      ("529_pollen", "pollen")]

rows = []
for key, label in DS:
    for arm, d, m in (("TSGP", "results_v7", "tsgp"),
                      ("standard GP", "results_v7", "stdgp"),
                      ("TSGP + control", "results_anneal_b", "tsgp"),
                      ("standard GP, no division", "results_nodiv", "stdgp")):
        rs = _runs(d, m).get(key)
        if not rs:
            continue
        md, mn, q1, q3, n = _stat(rs, "test_rmse")
        smd, _, sq1, sq3, _ = _stat(rs, "best_size")
        rows.append([label if arm == "TSGP" else "", arm, str(n),
                     f"{md:.4f}", f"{mn:.4f}", f"{q1:.4f} to {q3:.4f}",
                     f"{smd:.0f}", f"{sq1:.0f} to {sq3:.0f}"])
table_before(A, ["Dataset", "Arm", "n", "Median", "Mean", "IQR", "Size",
                 "Size IQR"], rows,
             [0.62, 1.22, 0.28, 0.66, 0.66, 1.06, 0.42, 0.78],
             left_cols=(1,))
caption_before(A, "Table C.1  Test RMSE and solution size on the regression "
                  "benchmarks, 30 runs per cell. The first two arms are the "
                  "replication of Chapter 6; the third is the extension of "
                  "Section 7.7 and is not equal-budget; the fourth is the "
                  "primitive-set comparison of Section 7.8.")

rows = []
for key, label in DS:
    for task, arm, d, m in (("median split", "TSGP", "results_clf_k1", "tsgp"),
                            ("median split", "standard GP", "results_clf",
                             "stdgp"),
                            ("middle band", "TSGP", "results_clfmid_k1",
                             "tsgp"),
                            ("middle band", "standard GP", "results_clfmid",
                             "stdgp")):
        rs = _runs(d, m).get(key)
        if not rs:
            continue
        md, mn, q1, q3, n = _stat(rs, "test_acc")
        auc = _stat(rs, "test_auc")[0]
        size = _stat(rs, "size")[0]
        maj = _stat(rs, "majority_baseline")[0]
        rows.append([label if task == "median split" and arm == "TSGP" else "",
                     task, arm, str(n), f"{md:.4f}",
                     f"{q1:.4f} to {q3:.4f}", f"{auc:.4f}", f"{size:.0f}",
                     f"{maj:.3f}"])
table_before(A, ["Dataset", "Task", "Arm", "n", "Accuracy", "IQR", "AUC",
                 "Size", "Majority"], rows,
             [0.6, 0.86, 0.94, 0.26, 0.7, 1.06, 0.58, 0.42, 0.72],
             left_cols=(1, 2))
caption_before(A, "Table C.2  Classification, both label constructions, at "
                  "equal budget and 30 runs per cell. Majority is the accuracy "
                  "of always predicting the more frequent class, and is the "
                  "level any usable model has to clear.")

# ----------------------------------------------------------------- glossary
# Every term is also defined at first use; the glossary is for the reader who
# meets one of them again forty pages later.
R = find_heading("References", level=1)
clear_section("Glossary of terms")
para_before(R, "Terms are defined at first use in the main text and are "
               "collected here for reference. The abbreviations used "
               "throughout are listed separately in the front matter.")
for term, definition in [
    ("Best-so-far reporting", "reporting the best solution found at any point "
     "in a run rather than the best in the final population. With "
     "generational replacement and no elitism the two differ, because the best "
     "individual is routinely destroyed before the run ends."),
    ("Bloat", "growth in program size across generations without a matching "
     "improvement in fitness."),
    ("Decision value", "the real-valued output of an expression read as a "
     "classifier: its sign gives the predicted class and its magnitude the "
     "confidence."),
    ("Desired semantic distance", "the distance requested of the learned "
     "operator when it is asked for an offspring, written SD_d. The paper sets "
     "it to 0.1; Chapter 8 rescales it to 2.0 for classification."),
    ("Double tournament", "a selection operator combining a fitness "
     "tournament with a parsimony tournament, used during data generation so "
     "that the pool collected is diverse without being dominated by large "
     "trees."),
    ("Ephemeral random constant", "a numeric leaf drawn, when a tree is built, "
     "from a fixed set of values: here the grid from -0.5 to 0.5 in steps of "
     "0.1."),
    ("Equal budget", "a comparison in which the methods being compared spend "
     "the same number of model evaluations per generation."),
    ("Geometric semantic operator", "a variation operator constructed so that "
     "the offspring provably lies at a controlled position between or near its "
     "parents in semantic space."),
    ("Grammar-constrained decoding", "masking, before each token is sampled, "
     "any token that would make a valid expression tree impossible, so that "
     "every generated sequence is syntactically valid and none has to be "
     "discarded and resampled."),
    ("Logistic loss", "the smooth classification objective log(1 + exp(-y "
     "f(x))), used as fitness in Chapter 8 because accuracy is piecewise "
     "constant and offers selection no gradient between boundary crossings."),
    ("Majority-class baseline", "the accuracy obtained by always predicting "
     "the more frequent class, and so the level any usable classifier has to "
     "clear."),
    ("Margin", "the quantity y f(x): positive when a prediction is correct, "
     "and larger the more confidently correct it is."),
    ("Median split", "the first label construction of Chapter 8: the target "
     "is thresholded at a cut chosen to divide the training half as evenly as "
     "the ties in these discrete targets allow."),
    ("Middle band", "the second label construction of Chapter 8: the positive "
     "class is the central third of the training distribution, which no linear "
     "model can represent."),
    ("Prefix notation", "writing a tree as a flat sequence in pre-order, so "
     "that add(x0, mul(x1, 0.3)) becomes [add, x0, mul, x1, 0.3]. This is what "
     "allows a tree to be treated as a sequence of tokens."),
    ("Protected division", "division that returns 1.0 when the absolute value "
     "of the denominator falls below 1e-6, rather than failing."),
    ("Ramped half-and-half", "an initialisation that mixes tree shapes and "
     "depths across the starting population."),
    ("Semantic distance", "the Euclidean distance between the semantics of "
     "two expressions."),
    ("Semantics", "the vector of outputs an expression produces on a fixed "
     "set of probe inputs. It describes what an expression does, as against "
     "how it is built."),
    ("Solution size", "the number of nodes in an expression tree."),
    ("Step-size control", "selecting, from several sampled offspring, the one "
     "whose semantic distance from the parent is closest to a target. The "
     "original paper lists this as future work; Section 7.7 supplies it "
     "externally."),
    ("Step-size floor", "the smallest parent-to-offspring semantic distance "
     "the learned operator can produce, about 0.7 here, which it does not fall "
     "below however much smaller a distance is requested."),
    ("Teacher forcing", "training a decoder on the true preceding tokens of "
     "the target rather than on the tokens it generated itself."),
    ("Tournament selection", "choosing the fittest of a randomly drawn subset "
     "of the population, of size five throughout this study."),
    ("Transfer control", "the arm in which the regression-trained operator is "
     "run on the classification benchmarks, which is what measures whether "
     "retraining on classification data was necessary."),
    ("Wilcoxon rank-sum test", "the non-parametric test of whether two "
     "independent samples differ in location, used at the 5% level throughout, "
     "as in the original paper."),
]:
    rich_before(R, [(term, True), (f"  {definition}", False)])

# -------------------------------------------------------------------- index
# Entries are written with a placeholder and filled from the rendered PDF in
# the same pass that fills the two front-matter lists, since neither the pages
# nor the terms can be known before the document is laid out.
INDEX_TERMS = [
    ("AdamW", r"AdamW"),
    ("batched sampler", r"batched (?:sampler|version)|batched decoder"),
    ("best-so-far reporting", r"best-so-far|best performing solution"),
    ("bloat", r"\bbloat"),
    ("classification, binary", r"binary classification"),
    ("DEAP", r"DEAP"),
    ("decision value", r"decision value"),
    ("Deep Symbolic Regression", r"Deep Symbolic Regression|\bDSR\b"),
    ("double tournament", r"double tournament"),
    ("ephemeral random constant", r"ephemeral random constant|\bERC\b"),
    ("equal budget", r"equal[- ]budget"),
    ("FAISS", r"FAISS"),
    ("geometric semantic GP", r"[Gg]eometric [Ss]emantic"),
    ("grammar-constrained decoding", r"[Gg]rammar-constrained|syntax control"),
    ("logistic loss", r"logistic loss"),
    ("logistic regression", r"[Ll]ogistic regression"),
    ("majority-class baseline", r"majority[- ](?:class )?baseline"),
    ("median split", r"median[- ]split|median split"),
    ("middle band", r"middle[- ]band"),
    ("nearest neighbours", r"nearest[- ]neighbour"),
    ("PMLB", r"PMLB|Penn Machine Learning"),
    ("prefix notation", r"prefix \(?(?:pre-order)?\)? ?notation|prefix "
                        r"notation"),
    ("protected division", r"[Pp]rotected division|protdiv"),
    ("ramped half-and-half", r"[Rr]amped [Hh]alf-and-[Hh]alf"),
    ("reproducibility", r"[Rr]eproducibilit"),
    ("RMSE", r"RMSE|root mean squared error"),
    ("semantic distance", r"semantic distance"),
    ("semantics", r"\bsemantics\b"),
    ("SLIM_GSGP", r"SLIM_GSGP"),
    ("solution size", r"solution size"),
    ("step-size control", r"step-size control|step selection|targeted step"),
    ("step-size floor", r"step-size floor|floor of|its floor"),
    ("teacher forcing", r"[Tt]eacher forcing"),
    ("tournament selection", r"[Tt]ournament selection"),
    ("training pairs", r"training pairs|five million pairs|semantic pair"),
    ("transformer", r"[Tt]ransformer"),
    ("Wilcoxon rank-sum test", r"Wilcoxon"),
]
clear_section("Index")
index_head = find_heading("Index", level=1)
para_before_index = doc.add_paragraph(
    "Page numbers refer to the chapters and appendices; terms are not indexed "
    "from the glossary, where each is defined in one place.", style=BODY)
index_head._p.addnext(para_before_index._p)
for term, _ in INDEX_TERMS:
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    p = doc.add_paragraph("", style=BODY)
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(6.4), WD_TAB_ALIGNMENT.RIGHT,
                              WD_TAB_LEADER.DOTS)
    # body spacing is set for prose; an index of one-line entries wants less
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    p.add_run(f"{term}\t{PLACEHOLDER_PAGE}")


# ============================================================ 2. renumber ==
def iter_body():
    """Paragraphs in document order, with their index."""
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)


def renumber_captions():
    chapter = 0
    in_chapters = False
    fig_n = tab_n = 0
    figures, tables = [], []
    for p in iter_body():
        st, txt = p.style.name, p.text.strip()
        if st == "Heading 1":
            if txt.lower().startswith(FIRST_CHAPTER.lower()):
                in_chapters = True
            if txt.lower().startswith(AFTER_LAST_CHAPTER.lower()):
                in_chapters = False
            if in_chapters:
                chapter += 1
                fig_n = tab_n = 0
            continue
        if st not in (FIGCAP, TABCAP) or not txt:
            continue
        kind = "Figure" if st == FIGCAP else "Table"
        # Appendix captions carry their own letters and sit outside the
        # chapter numbering, so they are collected for the lists but left as
        # written; only captions inside the numbered chapters are renumbered.
        if not in_chapters:
            m = re.match(r"((?:Figure|Table)\s+[A-Z]\.\d+)\s+(.*)", txt,
                         re.S)
            if m:
                (figures if kind == "Figure" else tables).append(
                    (m.group(1), m.group(2)))
            continue
        if kind == "Figure":
            fig_n += 1
            label = f"Figure {chapter}.{fig_n}"
        else:
            tab_n += 1
            label = f"Table {chapter}.{tab_n}"
        body_text = re.sub(rf"^(PLACEHOLDER|{kind}\s+\d+\.\d+)\s*", "", txt)
        for r in p.runs[1:]:
            r.text = ""
        p.runs[0].text = f"{label}  {body_text}"
        (figures if kind == "Figure" else tables).append((label, body_text))
    return figures, tables


figures, tables = renumber_captions()
print(f"renumbered {len(figures)} figures, {len(tables)} tables")


def clear_list(heading_text):
    """Empty a front-matter list, returning the node the entries go before."""
    from docx.text.paragraph import Paragraph
    head = find_heading(heading_text, level=1)
    node = head._p.getnext()
    while node is not None and node.tag == qn("w:p"):
        if Paragraph(node, doc).style.name.startswith("Heading"):
            break
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt
    return head




clear_list("List of Figures")
clear_list("List of Tables")


# ============================================================ 3. lists =====
# The page numbers in these lists were wrong in every earlier build, for two
# compounding reasons. They recorded the index of the page in the rendered
# PDF rather than the number printed on it, so every entry was out by the
# length of the front matter; and they were measured with the lists emptied,
# so filling them back in reflowed the document underneath the numbers that
# had just been measured.
#
# Both are fixed by writing the entries first with a placeholder the width of
# a real page number, so the document is already its final length when it is
# measured, then reading the number printed at the head of each page and
# writing it in without changing a single line. The measurement repeats until
# nothing moves, which is also what proves the numbers are stable. The index
# is filled from the same pass, for the same reason.
TMP_PDF = os.path.splitext(DST)[0] + "_tmp.pdf"


def rebuild_list(heading_text, entries):
    """Write the manual list under a front-matter heading (already cleared)."""
    head = find_heading(heading_text, level=1)
    anchor = head._p.getnext()
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    for label, body in entries:
        short = body.split(".")[0].strip()
        if len(short) > 62:                       # trim on a word boundary
            short = short[:62].rsplit(" ", 1)[0].rstrip(" ,;:") + " ..."
        p = doc.add_paragraph("", style=LIST_STYLE)
        # Without an explicit right-aligned tab stop the page number lands
        # wherever the default tab grid puts it, which for a long entry is
        # inside the text.
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.4), WD_TAB_ALIGNMENT.RIGHT,
                                  WD_TAB_LEADER.DOTS)
        pf.space_after = Pt(4)
        p.add_run(f"{label}  {short}\t{PLACEHOLDER_PAGE}")
        if anchor is not None:
            anchor.addprevious(p._p)
        else:
            head._p.addnext(p._p)


def refresh_and_render(docx_path, pdf_path):
    """Rebuild the table of contents in Word, then render the result.

    The contents page is a Word field. python-docx preserves its cached text
    rather than recomputing it, so on its own it would still show the chapter
    structure the source document had before Chapter 8 was inserted, with
    Threats at 8 and the Conclusion at 9. Word is the only thing that can
    rebuild it, and it has to happen before the pages are measured because a
    longer contents page moves everything after it.
    """
    for f in (pdf_path,):
        if os.path.exists(f):
            os.remove(f)
    ps = ("$w=New-Object -ComObject Word.Application;$w.Visible=$false;"
          "$w.DisplayAlerts=0;"
          f"$d=$w.Documents.Open('{docx_path}');"
          "foreach($t in $d.TablesOfContents){$t.Update()};"
          "$d.Fields.Update() | Out-Null;"
          "$d.Save();"
          f"$d.SaveAs([ref]'{pdf_path}',[ref]17);"
          "$d.Close(0);$w.Quit()")
    subprocess.run(["powershell", "-NoProfile", "-Command", ps],
                   capture_output=True)
    return os.path.exists(pdf_path)


def printed_pages(pdf_path, entries):
    """Where every caption lands, and every index term appears.

    Only pages numbered in arabic are considered, which excludes the front
    matter: the two lists repeat every caption label verbatim, and a plain
    label search would otherwise match the list entry rather than the caption.
    Captions are matched on their opening words as well as their label, so a
    sentence such as "as Table 6.1 shows" cannot claim the number instead.

    Index terms stop at the glossary, which defines all of them in one place
    and would otherwise appear against every entry, as would the index itself.
    """
    import pymupdf
    caption_pages, term_pages = {}, {}
    d = pymupdf.open(pdf_path)
    body = []
    for page in d:
        text = page.get_text()
        head = next((ln.strip() for ln in text.splitlines() if ln.strip()), "")
        if re.fullmatch(r"\d+", head):
            body.append((head, re.sub(r"\s+", " ", text)))
    d.close()

    for head, flat in body:
        for label, text in entries:
            if label in caption_pages:
                continue
            opening = re.sub(r"\s+", " ", text)[:25]
            if re.search(re.escape(label) + r"\s+" + re.escape(opening), flat):
                caption_pages[label] = head

    cut = next((i for i, (_, flat) in enumerate(body)
                if "Glossary of terms" in flat), len(body))
    for term, pattern in INDEX_TERMS:
        hits, seen = [], set()
        for head, flat in body[:cut]:
            if head not in seen and re.search(pattern, flat):
                seen.add(head)
                hits.append(head)
        if hits:
            # a term running through most of the report tells the reader
            # nothing by listing forty pages, and would wrap the entry
            term_pages[term] = ", ".join(hits[:8]) + \
                (" and passim" if len(hits) > 8 else "")
    return caption_pages, term_pages


def write_page_numbers(docx_path, pages, terms):
    """Set the measured numbers, changing no line's length.

    A caption entry is found by its label, an index entry by the term standing
    before the tab; nothing else in the document carries a tab, so the two
    cannot be confused with body text.
    """
    d2 = Document(docx_path)
    changed = 0
    for p in d2.paragraphs:
        if "\t" not in p.text:
            continue
        stem, _, current = p.text.rpartition("\t")
        m = re.match(r"((?:Figure|Table) [A-Z0-9]+\.\d+)\b", p.text)
        if m and m.group(1) in pages:
            want = pages[m.group(1)]
        elif stem in terms:
            want = terms[stem]
        else:
            continue
        if current == want:
            continue
        p.runs[0].text = f"{stem}\t{want}"
        for r in p.runs[1:]:
            r.text = ""
        changed += 1
    if changed:
        d2.save(docx_path)
    return changed


rebuild_list("List of Figures", figures)
rebuild_list("List of Tables", tables)
doc.save(DST)

entries = figures + tables
pages, terms = {}, {}
try:
    import pymupdf                                             # noqa: F401
    for attempt in range(1, 4):
        if not refresh_and_render(DST, TMP_PDF):
            print("  (could not render - page numbers left as placeholders)")
            break
        pages, terms = printed_pages(TMP_PDF, entries)
        os.remove(TMP_PDF)
        moved = write_page_numbers(DST, pages, terms)
        print(f"  pass {attempt}: located {len(pages)} of {len(entries)} "
              f"captions and {len(terms)} of {len(INDEX_TERMS)} index terms, "
              f"{moved} entries rewritten")
        if not moved:
            break
    else:
        print("  !! page numbers did not settle after three passes")
except ImportError:
    print("  (pymupdf not installed - page numbers left as placeholders)")

missing = [lab for lab, _ in entries if lab not in pages]
if missing:
    print(f"  !! not located in the rendered PDF: {', '.join(missing)}")
unfound = [t for t, _ in INDEX_TERMS if t not in terms]
if unfound:
    print(f"  !! index terms never matched: {', '.join(unfound)}")

print(f"\nwritten: {DST}")
print("\nfigures:")
for label, body in figures:
    print(f"  {label:<12} p{pages.get(label, '?'):<4} {body[:58]}")
print("tables:")
for label, body in tables:
    print(f"  {label:<12} p{pages.get(label, '?'):<4} {body[:58]}")
